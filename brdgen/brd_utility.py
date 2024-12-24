import os
import re
import unicodedata

import docx
from pypdf import PdfReader


class Utility:

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize input text."""
        print("Cleaning text...")
        text = re.sub(r"[\x00-\x1F\x7F-\x9F]", "", text)
        text = unicodedata.normalize("NFKD", text)
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from PDF or DOCX files."""
        print(f"Extracting text from {file_path}...")
        _, ext = os.path.splitext(file_path)

        try:
            if ext.lower() == ".pdf":
                with open(file_path, "rb") as file:
                    reader = PdfReader(file)
                    text = " ".join([page.extract_text() for page in reader.pages])
            elif ext.lower() in [".docx", ".doc"]:
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            # TODO add more file types if needed, like XLXS, CSV, etc.
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            return Utility.clean_text(text)

        except Exception as e:
            raise RuntimeError(f"Error extracting text from {file_path}: {str(e)}")

    @staticmethod
    def save_brd(brd_content: str, filename: str = "generated_brd.docx") -> str:
        # TODO: Add markdown to the brd_content
        print("Saving generated BRD")
        doc = docx.Document()
        doc.add_heading("Generated Business Requirements Document", level=1)
        doc.add_paragraph(brd_content)

        # Ensure directory exists
        os.makedirs("generated_brds", exist_ok=True)
        filepath = os.path.join("generated_brds", filename)
        doc.save(filepath)

        return filepath
